#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
应急预案规则提取器
使用API从Word文档中提取规则信息并转换为JSON格式
"""

import json
import re
import os
import hashlib
from pathlib import Path
from docx import Document
from typing import List, Dict, Any, Generator, Tuple
from openai import OpenAI
import tiktoken
import sys
from datetime import datetime
import math
SRC_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if SRC_DIR not in sys.path:
    sys.path.append(SRC_DIR)
from kg.config.config import API_CONFIG, PATH_CONFIG
from kg.utils.llm_client import LLMClientFactory

class RuleExtractor:
    def __init__(self, model: str = "deepseek-r1", max_tokens: int = 2000, temperature: float = 0.1):
        self.rules = []
        self.rule_counter = 1
        self.seen_rule_texts = set()  # 用于去重
        # API配置
        self.api_key = os.getenv("DASHSCOPE_API_KEY", "")
        if not self.api_key:
            raise RuntimeError("missing DASHSCOPE_API_KEY")
        self.model = model
        self.max_tokens = max_tokens
        self.temperature = temperature
        self.client = OpenAI(
            api_key=self.api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )
        self.zhipu_client = LLMClientFactory.create_client(
            provider="zhipu",
            api_key=API_CONFIG.get("zhipu_api_key", ""),
            model="embedding-3",
            timeout=30,
            max_retries=3
        )
        self.embedding_output_dir = os.path.join(PATH_CONFIG.get("output_dir") or "", "embeddings")
        os.makedirs(self.embedding_output_dir, exist_ok=True)
    
    def num_tokens_from_string(self, string: str, encoding_name: str = "cl100k_base") -> int:
        """
        计算字符串的token数量
        """
        try:
            encoding = tiktoken.get_encoding(encoding_name)
            num_tokens = len(encoding.encode(string))
            return num_tokens
        except Exception:
            # 如果tiktoken不可用，使用简单的字符数估算（1个token约等于4个字符）
            return len(string) // 4
    
    def concat_img(self, img1, img2):
        """
        合并图像（简化版本）
        """
        if img1 and not img2:
            return img1
        if img2 and not img1:
            return img2
        if img1 and img2:
            return img1  # 简化处理，返回第一个图像
        return None
    
    def naive_merge_docx(self, sections, chunk_token_num=1024, delimiter="\n。；！？"):
        """
        RAGflow的docx分块方法
        
        基于token数量和语义边界对文档进行智能分块。
        该方法确保每个分块的token数量不超过指定阈值，
        同时尽量保持语义的完整性。
        
        Args:
            sections (list): 文档段落列表，每个元素为(文本, 图像)元组
            chunk_token_num (int): 每个分块的最大token数量，默认1024
            delimiter (str): 分块边界标识符，默认为中文标点符号
            
        Returns:
            tuple: (分块文本列表, 分块图像列表)
        """
        if not sections:
            return [], []
        cks: List[str] = []
        images: List[Any] = []

        cur_text = ""
        cur_img = None
        cur_tokens = 0

        last_para_text = None
        last_para_img = None

        for sec, image in sections:
            t = str(sec or "")
            if not t:
                continue
            tnum = self.num_tokens_from_string(t)

            if cur_text and (cur_tokens + tnum > chunk_token_num):
                cks.append(cur_text)
                images.append(cur_img)

                cur_text = ""
                cur_img = None
                cur_tokens = 0

                if last_para_text:
                    cur_text += last_para_text
                    cur_img = self.concat_img(cur_img, last_para_img)
                    cur_tokens = self.num_tokens_from_string(last_para_text)

            cur_text += t
            cur_img = self.concat_img(cur_img, image)
            cur_tokens += tnum

            last_para_text = t
            last_para_img = image

        if cur_text:
            cks.append(cur_text)
            images.append(cur_img)

        return cks, images

    def _normalize_semantic_element_key(self, k: Any) -> str:
        s = str(k or "").strip()
        if not s:
            return ""
        s_low = s.lower()
        if s_low in {"s", "subject", "主体"}:
            return "Subject"
        if s_low in {"a", "action", "动作"}:
            return "Action"
        if s_low in {"c", "condition", "条件"}:
            return "Condition"
        if s_low in {"o", "object", "对象", "范围"}:
            return "Object"
        if s_low in {"t", "time", "时限", "时间"}:
            return "Time"
        if s_low in {"subject"}:
            return "Subject"
        if s_low in {"action"}:
            return "Action"
        if s_low in {"condition"}:
            return "Condition"
        if s_low in {"object"}:
            return "Object"
        if s_low in {"time"}:
            return "Time"
        return ""

    def _extract_semantic_elements(self, rule_data: Dict[str, Any], rule_text: str) -> set:
        elements = set()

        if isinstance(rule_data, dict):
            for key in ("Subject", "Action", "Condition", "Object", "Time", "S", "A", "C", "O", "T", "主体", "动作", "条件", "对象", "范围", "时限", "时间"):
                if key not in rule_data:
                    continue
                v = rule_data.get(key)
                if v is None:
                    continue
                if isinstance(v, (list, tuple, set, dict)):
                    if len(v) == 0:
                        continue
                else:
                    if not str(v).strip():
                        continue
                norm = self._normalize_semantic_element_key(key)
                if norm:
                    elements.add(norm)

            extra = rule_data.get("elements") or rule_data.get("Elements") or None
            if isinstance(extra, (list, tuple, set)):
                for item in extra:
                    norm = self._normalize_semantic_element_key(item)
                    if norm:
                        elements.add(norm)
            elif isinstance(extra, dict):
                for k2, v2 in extra.items():
                    if v2 is None:
                        continue
                    if isinstance(v2, (list, tuple, set, dict)):
                        if len(v2) == 0:
                            continue
                    else:
                        if not str(v2).strip():
                            continue
                    norm = self._normalize_semantic_element_key(k2)
                    if norm:
                        elements.add(norm)

        txt = str(rule_text or "").strip()
        if txt:
            action_markers = ["应当", "必须", "不得", "禁止", "应及时", "应立即", "需要", "要求", "负责", "建立", "启动", "报告", "发布", "处置", "组织", "统筹", "协调", "完善", "落实"]
            if any(m in txt for m in action_markers):
                elements.add("Action")

            if re.search(r"(当|若|如果|如遇|如发生|发生|出现|遇到|在.+?时|在.+?期间)", txt):
                elements.add("Condition")

            if re.search(r"(\d+\s*(分钟|小时|天|日|月)\s*(内|以内|之内)?|立即|及时|尽快|第一时间)", txt):
                elements.add("Time")

            m = re.search(r"(应当|必须|不得|禁止|需要|要求|负责|组织|建立|启动|报告|发布|处置|统筹|协调|完善|落实)", txt)
            if m and m.start() >= 2:
                elements.add("Subject")

            if re.search(r"(对|向|将|把)\S{1,40}", txt):
                elements.add("Object")

        return elements

    def _passes_rule_constraints(self, elements: set) -> bool:
        if not elements or len(elements) < 2:
            return False
        return ("Action" in elements) or ("Condition" in elements)
    
    def __embed_texts(self, texts: List[str], dimensions: int = None) -> List[List[float]]:
        return self.zhipu_client.embed(texts, dimensions=dimensions)
    
    def __save_jsonl(self, records: List[Dict[str, Any]], out_path: str):
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as f:
            for rec in records:
                f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    
    def embed_and_save_chunks(self, document_name: str, chunk_texts: List[str], dim: int = None) -> str:
        if not chunk_texts:
            return ""
        ts = datetime.utcnow().isoformat()
        records = []
        batch_size = 64
        for start in range(0, len(chunk_texts), batch_size):
            batch = chunk_texts[start:start+batch_size]
            vecs = self.__embed_texts(batch, dimensions=dim)
            for i, t in enumerate(batch):
                v = vecs[i] if i < len(vecs) else []
                records.append({
                    "chunk_text": t,
                    "embedding": v,
                    "model": "embedding-3",
                    "dim": len(v),
                    "created_at": ts,
                    "norm_l2": math.sqrt(sum(x*x for x in v)) if v else 0.0,
                    "source": {
                        "document": document_name,
                        "chunk_index": start + i
                    }
                })
        out_path = os.path.join(self.embedding_output_dir, f"{document_name}_chunk_embeddings.jsonl")
        self.__save_jsonl(records, out_path)
        return out_path
    
    def embed_and_save_rules(self, rules: List[Dict[str, Any]], dim: int = None) -> str:
        if not rules:
            return ""
        ts = datetime.utcnow().isoformat()
        records = []
        texts = [r.get("rule_text", "") for r in rules]
        batch_size = 64
        for start in range(0, len(texts), batch_size):
            batch = texts[start:start+batch_size]
            vecs = self.__embed_texts(batch, dimensions=dim)
            for i, r in enumerate(rules[start:start+batch_size]):
                v = vecs[i] if i < len(vecs) else []
                records.append({
                    "rule_id": r.get("rule_id"),
                    "rule_text": r.get("rule_text"),
                    "embedding": v,
                    "model": "embedding-3",
                    "dim": len(v),
                    "created_at": ts,
                    "norm_l2": math.sqrt(sum(x*x for x in v)) if v else 0.0,
                    "source": {
                        "rule_name": r.get("rule_name"),
                        "rule_source": r.get("rule_source"),
                        "rule_tags": r.get("rule_tags")
                    }
                })
        document_name = rules[0].get("rule_source", "unknown")
        out_path = os.path.join(self.embedding_output_dir, f"{document_name}_rule_embeddings.jsonl")
        self.__save_jsonl(records, out_path)
        return out_path
    
    def __cosine(self, a: List[float], b: List[float]) -> float:
        if not a or not b:
            return 0.0
        if len(a) != len(b):
            return 0.0
        dot = sum(x*y for x, y in zip(a, b))
        na = math.sqrt(sum(x*x for x in a))
        nb = math.sqrt(sum(y*y for y in b))
        if na == 0 or nb == 0:
            return 0.0
        return dot / (na * nb)
    
    def semantic_search(self, document_name: str, query_text: str, top_k: int = 5, rule_threshold: float = 0.6) -> Dict[str, Any]:
        qv = self.__embed_texts([query_text])[0]
        rule_path = os.path.join(self.embedding_output_dir, f"{document_name}_rule_embeddings.jsonl")
        chunk_path = os.path.join(self.embedding_output_dir, f"{document_name}_chunk_embeddings.jsonl")
        rule_results = []
        if os.path.exists(rule_path):
            with open(rule_path, "r", encoding="utf-8") as f:
                for line in f:
                    rec = json.loads(line)
                    sim = self.__cosine(qv, rec.get("embedding", []))
                    rule_results.append({
                        "type": "rule",
                        "id": rec.get("rule_id"),
                        "text": rec.get("rule_text"),
                        "similarity": sim
                    })
            rule_results.sort(key=lambda x: x["similarity"], reverse=True)
        if rule_results and rule_results[0]["similarity"] >= rule_threshold:
            return {"query": query_text, "top_k": top_k, "results": rule_results[:top_k]}
        chunk_results = []
        if os.path.exists(chunk_path):
            with open(chunk_path, "r", encoding="utf-8") as f:
                for line in f:
                    rec = json.loads(line)
                    sim = self.__cosine(qv, rec.get("embedding", []))
                    chunk_results.append({
                        "type": "chunk",
                        "id": rec.get("chunk_index"),
                        "text": rec.get("chunk_text"),
                        "similarity": sim
                    })
            chunk_results.sort(key=lambda x: x["similarity"], reverse=True)
        return {"query": query_text, "top_k": top_k, "results": (rule_results[:top_k] if rule_results else []) or chunk_results[:top_k]}
    
    def extract_docx_sections_for_ragflow(self, doc):
        """
        提取docx文档内容，格式适配RAGflow分块函数
        
        从Word文档中提取段落和表格内容，转换为RAGflow
        分块函数所需的格式：(文本内容, 图像)元组列表。
        
        Args:
            doc: DocxDocument对象
            
        Returns:
            List[Tuple]: 段落列表，每个元素为(文本, 图像)元组
        """
        sections = []
        
        # 提取段落
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # 只保留非空段落
                sections.append((text, None))
        
        # 提取表格内容
        for table in doc.tables:
            table_text = ""
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text += " | ".join(row_text) + "\n"
            
            if table_text.strip():
                sections.append((table_text.strip(), None))
        
        return sections
        
    def extract_from_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        使用智能分块和API从Word文档中提取规则信息
        """
        try:
            doc = Document(docx_path)
            document_name = Path(docx_path).stem
            
            # 使用智能分块系统
            print("正在提取文档结构...")
            sections = self.extract_docx_sections_for_ragflow(doc)
            print(f"提取到 {len(sections)} 个文档段落")
            
            # 使用RAGflow智能分块
            print("正在进行智能分块...")
            chunk_texts, chunk_images = self.naive_merge_docx(sections, chunk_token_num=1024)
            print(f"智能分块完成，共生成 {len(chunk_texts)} 个分块")
            self.embed_and_save_chunks(document_name, chunk_texts, dim=512)
            
            # 使用API提取规则
            rules = self._extract_rules_with_api_chunks(chunk_texts, document_name)
            
            # 去重处理
            unique_rules = self._deduplicate_rules(rules)
            print(f"去重后保留 {len(unique_rules)} 条规则")
            
            self.rules.extend(unique_rules)
            try:
                self.embed_and_save_rules(unique_rules, dim=512)
            except Exception as e:
                print(f"条例向量生成失败：{e}，将仅保存规则JSON")
                
            return self.rules
            
        except Exception as e:
            print(f"处理文档时出错: {e}")
            return []
    
    def _extract_rules_with_api_chunks(self, chunk_texts: List[str], document_name: str) -> List[Dict[str, Any]]:
        """
        使用API从智能分块中提取规则信息
        """
        rules = []
        
        print(f"开始处理 {len(chunk_texts)} 个智能分块")
        
        for i, chunk_text in enumerate(chunk_texts):
            if not chunk_text.strip():
                continue
                
            print(f"正在处理第 {i+1}/{len(chunk_texts)} 个分块...")
            
            try:
                prompt = f"""
请从以下文档分块中抽取“应急预案条例”（具有约束力的规范性条款）。

【重要说明】
- 若该分块没有明确的条例，返回 []
- 只提取具有规范性/指导性/约束性的句子，排除纯描述或背景介绍
- 输出的条例必须是语法完整的单句，能直接用于执行或监督

【判定标准】
同时满足以下至少两项：
1) 含有规范性词：应当、必须、不得、禁止、应及时、应立即、需要、要求、负责、建立、启动、报告、发布、处置、组织、统筹、协调、完善、落实
2) 明确规定职责、流程、条件、时限、标准、级别（如 I/II/III/IV 级响应）
3) 具备可操作性与可监督性，可直接据此开展工作或考核

【结构完整性】
- 条例需覆盖 主体-动作-对象/范围-条件/时限 中的至少两个要素；优先补全缺失主语
- 若分块来自表格或要点列表，请意译为完整句式，不保留原编号
- 避免仅返回“负责…/组织…”等缺主语或条件的信息片段

【输出字段】
按数组返回，每条包含：
- rule_name：简洁的条例名称（不含编号）
- rule_text：完整条例文本（包含主语+动作，必要时补足条件/时限/级别）
- Subject/Action/Condition/Object/Time：从条例中抽取的语义要素（缺失可为空字符串）
- rule_tags：选择 2–4 个标签进行分类（从下列集合中选取）：
  [响应等级, 预警, 信息报送, 指挥, 职责, 措施, 监测, 救援, 物资, 人员, 培训, 演练, 评估, 管理, 协调, 通信, 决策, 恢复, 撤离, 值守]

【文档分块】
{chunk_text}

【返回格式】
仅返回 JSON 数组；无内容则返回 []：
[
  {{
    "rule_name": "名称",
    "rule_text": "完整条例文本（包含主语与动作，必要时补足条件/时限）",
    "Subject": "主体",
    "Action": "动作",
    "Condition": "条件",
    "Object": "对象/范围",
    "Time": "时限/时间",
    "rule_tags": ["标签1", "标签2"]
  }}
]

【示例】
- 错误示例："负责组织全市汛情对外发布。"（缺主语）
- 正确示例："市防汛抗旱指挥部办公室应及时组织全市汛情对外发布。"
"""
                
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "只返回JSON数组，不要输出代码块或解释。如果没有规则，返回[]。严格遵守合规要求，避免输出不适宜内容。"},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=self.max_tokens,
                    temperature=self.temperature
                )
                
                # 解析API响应
                api_response = response.choices[0].message.content.strip()
                
                # 尝试解析JSON
                try:
                    chunk_rules = None
                    try:
                        parsed = json.loads(api_response)
                        if isinstance(parsed, list):
                            chunk_rules = parsed
                        elif isinstance(parsed, dict):
                            for key in ("rules", "data", "output"):
                                val = parsed.get(key)
                                if isinstance(val, list):
                                    chunk_rules = val
                                    break
                    except Exception:
                        code_blocks = re.findall(r"```json\s*([\s\S]*?)```", api_response)
                        for block in code_blocks:
                            try:
                                chunk_rules = json.loads(block)
                                break
                            except Exception:
                                pass
                        if chunk_rules is None:
                            arrays = re.findall(r"\[[\s\S]*?\]", api_response)
                            for arr in arrays:
                                try:
                                    chunk_rules = json.loads(arr)
                                    break
                                except Exception:
                                    pass
                    if not isinstance(chunk_rules, list):
                        chunk_rules = []
                    if chunk_rules is not None:
                        
                        # 为每条规则添加ID和来源
                        kept = 0
                        for rule_data in chunk_rules:
                            if isinstance(rule_data, dict) and 'rule_text' in rule_data and rule_data['rule_text'].strip():
                                elements = self._extract_semantic_elements(rule_data, rule_data.get("rule_text", ""))
                                if not self._passes_rule_constraints(elements):
                                    continue
                                rule = {
                                    "rule_id": str(self.rule_counter),
                                    "rule_name": rule_data.get('rule_name', '未命名规则'),
                                    "rule_text": rule_data.get('rule_text', '').strip(),
                                    "elements": sorted(elements),
                                    "rule_tags": rule_data.get('rule_tags', ['应急管理']),
                                    "rule_source": document_name
                                }
                                rules.append(rule)
                                self.rule_counter += 1
                                kept += 1
                        
                        if chunk_rules:
                            print(f"第 {i+1} 个分块提取到 {kept} 条规则")
                        else:
                            print(f"第 {i+1} 个分块未发现规则条例")
                    else:
                        print(f"第 {i+1} 个分块未找到有效的JSON格式")
                        
                except json.JSONDecodeError as e:
                    print(f"第 {i+1} 个分块JSON解析失败: {e}")
                    print(f"API响应前200字符: {api_response[:200]}...")
                    
            except Exception as e:
                err_text = str(e)
                if "data_inspection_failed" in err_text.lower() or "inappropriate content" in err_text.lower():
                    try:
                        print(f"第 {i+1} 个分块触发内容安全拦截，尝试使用更严格的JSON输出格式重试...")
                        safe_messages = [
                            {"role": "system", "content": "只返回JSON数组，确保内容合法合规，不包含不适宜内容。如果没有规则，返回[]。"},
                            {"role": "user", "content": prompt}
                        ]
                        response = self.client.chat.completions.create(
                            model=self.model,
                            messages=safe_messages,
                            max_tokens=self.max_tokens,
                            temperature=self.temperature
                        )
                        api_response = response.choices[0].message.content.strip()
                        try:
                            parsed = json.loads(api_response)
                            if isinstance(parsed, list):
                                chunk_rules = parsed
                            elif isinstance(parsed, dict):
                                for key in ("rules", "data", "output"):
                                    val = parsed.get(key)
                                    if isinstance(val, list):
                                        chunk_rules = val
                                        break
                            else:
                                chunk_rules = []
                        except Exception:
                            chunk_rules = []
                        if chunk_rules:
                            kept = 0
                            for rule_data in chunk_rules:
                                if isinstance(rule_data, dict) and 'rule_text' in rule_data and rule_data['rule_text'].strip():
                                    elements = self._extract_semantic_elements(rule_data, rule_data.get("rule_text", ""))
                                    if not self._passes_rule_constraints(elements):
                                        continue
                                    rule = {
                                        "rule_id": str(self.rule_counter),
                                        "rule_name": rule_data.get('rule_name', '未命名规则'),
                                        "rule_text": rule_data.get('rule_text', '').strip(),
                                        "elements": sorted(elements),
                                        "rule_tags": rule_data.get('rule_tags', ['应急管理']),
                                        "rule_source": document_name
                                    }
                                    rules.append(rule)
                                    self.rule_counter += 1
                                    kept += 1
                            print(f"第 {i+1} 个分块重试成功，提取到 {kept} 条规则")
                            continue
                    except Exception as re:
                        print(f"第 {i+1} 个分块重试仍失败: {re}")
                print(f"第 {i+1} 个分块API调用失败: {e}")
                continue
        
        return rules
    
    def _deduplicate_rules(self, rules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        去除重复的规则文本
        """
        unique_rules = []
        seen_texts = set()
        
        for rule in rules:
            rule_text = rule.get('rule_text', '').strip()
            if rule_text and rule_text not in seen_texts:
                seen_texts.add(rule_text)
                unique_rules.append(rule)
            else:
                print(f"发现重复规则，已跳过: {rule_text[:50]}...")
        
        return unique_rules
    

    
    def save_to_json(self, output_path: str):
        """
        保存规则到JSON文件
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.rules, f, ensure_ascii=False, indent=4)
            print(f"成功保存 {len(self.rules)} 条规则到 {output_path}")
        except Exception as e:
            print(f"保存文件时出错: {e}")

def main():
    """
    批量处理数据集 下的应急预案 Word 文档，逐个提取条例并输出相邻 JSON 文件。
    支持 .docx，.doc 将尝试自动转换为 .docx（需本机安装 Microsoft Word）。
    """
    input_dir = r""
    print("=== 应急预案规则提取器（批处理） ===")
    print(f"输入目录: {input_dir}")

    if not os.path.isdir(input_dir):
        print(f"输入目录不存在：{input_dir}")
        return

    def convert_doc_to_docx(doc_path: str) -> str:
        """尝试将 .doc 转为 .docx，返回转换后的路径；失败则返回空字符串。"""
        try:
            import win32com.client as win32
        except Exception:
            print(f"未检测到 win32com，跳过 .doc 转换：{doc_path}")
            return ""
        try:
            app = win32.Dispatch('Word.Application')
            app.Visible = False
            doc = app.Documents.Open(doc_path)
            out_path = doc_path + "x" if not doc_path.lower().endswith(".docx") else doc_path
            doc.SaveAs(out_path, FileFormat=16)
            doc.Close(False)
            app.Quit()
            print(f"已转换为 DOCX：{out_path}")
            return out_path
        except Exception as e:
            print(f".doc 转换失败，跳过：{doc_path}，错误：{e}")
            return ""

    files = [f for f in os.listdir(input_dir) if f.lower().endswith((".docx", ".doc"))]
    if not files:
        print("目录下未发现 .docx/.doc 文件")
        return

    total_files = len(files)
    print(f"发现 {total_files} 个文档，开始逐个提取…")

    processed = 0
    for idx, fname in enumerate(files, 1):
        fpath = os.path.join(input_dir, fname)
        # 准备 docx 路径
        ext = os.path.splitext(fname)[1].lower()
        if ext == ".docx":
            docx_path = fpath
        else:
            docx_path = convert_doc_to_docx(fpath)
            if not docx_path:
                continue

        print(f"\n[{idx}/{total_files}] 处理：{fname}")
        # 每个文件单独创建提取器，保证计数与集合独立
        try:
            extractor = RuleExtractor()
            print("API客户端初始化成功")
        except Exception as e:
            print(f"API客户端初始化失败：{e}")
            continue

        try:
            rules = extractor.extract_from_docx(docx_path)
            if rules:
                out_json = os.path.splitext(docx_path if ext == ".docx" else fpath)[0] + ".json"
                extractor.save_to_json(out_json)
                print(f"完成：{fname} -> {os.path.basename(out_json)}，规则数：{len(rules)}")
                processed += 1
            else:
                print(f"未提取到规则：{fname}")
        except Exception as e:
            print(f"提取失败：{fname}，错误：{e}")

    print(f"\n=== 批处理完成：成功处理 {processed}/{total_files} 个文档 ===")
    print("\n前3条规则示例:")
    for i, rule in enumerate(rules[:3], 1):
            print(f"\n--- 规则 {i} ---")
            print(f"ID: {rule['rule_id']}")
            print(f"名称: {rule['rule_name']}")
            print(f"内容: {rule['rule_text'][:150]}{'...' if len(rule['rule_text']) > 150 else ''}")
            print(f"标签: {', '.join(rule['rule_tags'])}")
            print(f"来源: {rule['rule_source']}")
    else:
        print("\n未提取到任何规则，请检查文档内容或API配置")

if __name__ == "__main__":
    main()
